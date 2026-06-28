"""OCR routing and text extraction.

Routing (per ARCHITECTURE.md):
  * PDF  → probe for a usable text layer with PyMuPDF; if present, read it
           directly; otherwise rasterize each page and OCR via Google Vision.
  * DOCX → read text directly with python-docx.
  * image → Google Vision.

Output is normalized to an `OcrResult`. Vision paths keep block-level bounding
boxes (provenance depends on them and it is painful to retrofit). Results are
cached by file hash — identical bytes are never OCR'd twice, across accounts.
"""

from __future__ import annotations

import json
import time
import uuid
from collections import Counter
from dataclasses import asdict, dataclass, field
from functools import partial
from pathlib import Path

import fitz  # PyMuPDF
from langdetect import DetectorFactory, LangDetectException, detect

from app.core.concurrency import map_bounded
from app.core.config import get_settings
from app.core.retry import with_retry
from app.db.models import Document
from app.db.session import SessionLocal
from app.services.events import record_event
from app.services.storage import get_storage_root

DetectorFactory.seed = 0  # deterministic language detection

# --- MIME handling ---------------------------------------------------------
PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
IMAGE_MIMES = frozenset({"image/png", "image/jpeg"})
ALLOWED_MIME_TYPES = frozenset({PDF_MIME, DOCX_MIME, *IMAGE_MIMES})

_EXT_BY_MIME = {
    PDF_MIME: ".pdf",
    DOCX_MIME: ".docx",
    "image/png": ".png",
    "image/jpeg": ".jpg",
}

# --- OCR engine names (mirror the ocr_engine enum) -------------------------
ENGINE_PDF_TEXT = "pdf_text_layer"
ENGINE_VISION = "google_vision"
ENGINE_DOCX = "docx"

# --- text-layer probe thresholds -------------------------------------------
# A PDF "has a usable text layer" only if it carries enough real characters;
# scanned PDFs typically yield near-zero extractable text.
_MIN_TOTAL_CHARS = 100
_MIN_CHARS_PER_PAGE = 20

# DPI used when rasterizing PDF pages for Vision fallback.
_RASTER_DPI = 200


def extension_for(mime_type: str, filename: str | None = None) -> str:
    """Return the storage extension for a MIME type (falls back to filename)."""
    if mime_type in _EXT_BY_MIME:
        return _EXT_BY_MIME[mime_type]
    if filename and "." in filename:
        return Path(filename).suffix.lower()
    return ".bin"


@dataclass
class OcrBlock:
    text: str
    bbox: list[list[float]]  # polygon vertices [[x, y], ...]


@dataclass
class OcrPage:
    page: int
    text: str
    blocks: list[OcrBlock] = field(default_factory=list)


@dataclass
class OcrResult:
    engine: str
    page_count: int
    language: str | None
    text: str
    pages: list[OcrPage] = field(default_factory=list)
    failed_pages: list[int] = field(default_factory=list)  # pages Vision dropped

    def to_cache(self) -> dict:
        return asdict(self)

    @classmethod
    def from_cache(cls, data: dict) -> "OcrResult":
        pages = [
            OcrPage(
                page=p["page"],
                text=p["text"],
                blocks=[OcrBlock(**b) for b in p.get("blocks", [])],
            )
            for p in data.get("pages", [])
        ]
        return cls(
            engine=data["engine"],
            page_count=data["page_count"],
            language=data.get("language"),
            text=data["text"],
            pages=pages,
            failed_pages=data.get("failed_pages", []),
        )


# --- engine selection (pure, unit-tested) ----------------------------------
def choose_engine(mime_type: str, has_text_layer: bool) -> str:
    """Decide which OCR engine handles `mime_type`."""
    if mime_type == PDF_MIME:
        return ENGINE_PDF_TEXT if has_text_layer else ENGINE_VISION
    if mime_type == DOCX_MIME:
        return ENGINE_DOCX
    if mime_type in IMAGE_MIMES:
        return ENGINE_VISION
    raise ValueError(f"Unsupported MIME type for OCR routing: {mime_type!r}")


# --- PDF text-layer probe --------------------------------------------------
def probe_pdf_text_layer(path: str | Path) -> tuple[list[str], int, bool]:
    """Extract per-page text from a PDF and judge whether the layer is usable.

    Returns ``(page_texts, page_count, has_usable_layer)``.
    """
    page_texts: list[str] = []
    with fitz.open(path) as doc:
        page_count = doc.page_count
        for page in doc:
            page_texts.append(page.get_text("text"))

    total_chars = sum(len(t.strip()) for t in page_texts)
    has_layer = (
        page_count > 0
        and total_chars >= _MIN_TOTAL_CHARS
        and (total_chars / page_count) >= _MIN_CHARS_PER_PAGE
    )
    return page_texts, page_count, has_layer


# --- docx ------------------------------------------------------------------
def extract_docx(path: str | Path) -> str:
    """Read visible paragraph text from a .docx file."""
    import docx  # local import keeps module import cheap

    document = docx.Document(str(path))
    return "\n".join(p.text for p in document.paragraphs if p.text.strip())


# --- language --------------------------------------------------------------
def detect_language(text: str) -> str | None:
    """Best-effort ISO-639-1 language code, or None for too-short/unknown text."""
    sample = text.strip()
    if len(sample) < 20:
        return None
    try:
        return detect(sample)
    except LangDetectException:
        return None


# --- Google Vision ---------------------------------------------------------
_vision_singleton = None


def _vision_client():
    global _vision_singleton
    if _vision_singleton is None:
        from google.cloud import vision

        _vision_singleton = vision.ImageAnnotatorClient()
    return _vision_singleton


def _blocks_from_annotation(annotation) -> tuple[str, list[OcrBlock], list[str]]:
    """Flatten a Vision full_text_annotation into text + block bboxes + languages."""
    blocks: list[OcrBlock] = []
    languages: list[str] = []
    for page in annotation.pages:
        for lang in getattr(page.property, "detected_languages", []) or []:
            if lang.language_code:
                languages.append(lang.language_code)
        for block in page.blocks:
            words: list[str] = []
            for paragraph in block.paragraphs:
                for word in paragraph.words:
                    words.append("".join(sym.text for sym in word.symbols))
            block_text = " ".join(words).strip()
            if not block_text:
                continue
            bbox = [[v.x, v.y] for v in block.bounding_box.vertices]
            blocks.append(OcrBlock(text=block_text, bbox=bbox))
    full_text = annotation.text or "\n".join(b.text for b in blocks)
    return full_text, blocks, languages


def _is_transient_vision(exc: Exception) -> bool:
    """True for transient Google Vision errors worth retrying (never 4xx/auth)."""
    try:
        from google.api_core import exceptions as gexc
    except ImportError:  # pragma: no cover - google-cloud-vision is a hard dependency
        return False
    return isinstance(
        exc,
        (
            gexc.ServiceUnavailable,    # 503
            gexc.TooManyRequests,       # 429
            gexc.DeadlineExceeded,      # 504 / timeout
            gexc.InternalServerError,   # 500
            gexc.GatewayTimeout,        # 504
        ),
    )


def _vision_ocr_with_retry(content: bytes) -> tuple[str, list[OcrBlock], list[str]]:
    """Vision OCR of one image, retrying transient failures."""
    settings = get_settings()
    return with_retry(
        partial(_vision_ocr_image_bytes, content),
        attempts=settings.retry_max_attempts,
        base_delay=settings.retry_base_delay,
        is_retryable=_is_transient_vision,
    )


def _vision_ocr_image_bytes(content: bytes) -> tuple[str, list[OcrBlock], list[str]]:
    from google.cloud import vision

    response = _vision_client().document_text_detection(image=vision.Image(content=content))
    if response.error.message:
        raise RuntimeError(f"Vision API error: {response.error.message}")
    return _blocks_from_annotation(response.full_text_annotation)


def ocr_image_via_vision(path: str | Path) -> OcrResult:
    content = Path(path).read_bytes()
    text, blocks, languages = _vision_ocr_with_retry(content)
    language = Counter(languages).most_common(1)[0][0] if languages else detect_language(text)
    return OcrResult(
        engine=ENGINE_VISION,
        page_count=1,
        language=language,
        text=text,
        pages=[OcrPage(page=1, text=text, blocks=blocks)],
    )


def ocr_pdf_via_vision(path: str | Path) -> OcrResult:
    """Rasterize each PDF page (PyMuPDF) and OCR it with Vision.

    Partial tolerance: a page that keeps failing after retries is recorded in
    `failed_pages` and left with empty text rather than failing the whole
    document; only an all-pages failure raises.
    """
    # Rasterize serially — PyMuPDF is not thread-safe on a shared document — then
    # OCR the page images in parallel (the network phase).
    zoom = _RASTER_DPI / 72.0
    matrix = fitz.Matrix(zoom, zoom)
    with fitz.open(path) as doc:
        page_count = doc.page_count
        rasters = [
            (index, page.get_pixmap(matrix=matrix).tobytes("png"))
            for index, page in enumerate(doc, start=1)
        ]

    def _ocr_page(item: tuple[int, bytes]) -> tuple[int, str | None, list[OcrBlock], list[str]]:
        index, content = item
        try:
            text, blocks, langs = _vision_ocr_with_retry(content)
        except Exception:  # noqa: BLE001 — tolerate a bad page, record it
            return index, None, [], []
        return index, text, blocks, langs

    outcomes = map_bounded(
        _ocr_page, rasters, max_workers=get_settings().max_parallel_calls
    )

    pages: list[OcrPage] = []
    languages: list[str] = []
    failed_pages: list[int] = []
    for index, text, blocks, langs in outcomes:  # input (page) order preserved
        if text is None:
            failed_pages.append(index)
            pages.append(OcrPage(page=index, text="", blocks=[]))
        else:
            languages.extend(langs)
            pages.append(OcrPage(page=index, text=text, blocks=blocks))

    if failed_pages and len(failed_pages) == page_count:
        raise RuntimeError(f"Vision OCR failed for all {page_count} page(s)")

    full_text = "\n\n".join(p.text for p in pages).strip()
    language = Counter(languages).most_common(1)[0][0] if languages else detect_language(full_text)
    return OcrResult(
        engine=ENGINE_VISION,
        page_count=page_count,
        language=language,
        text=full_text,
        pages=pages,
        failed_pages=failed_pages,
    )


# --- OCR cache (keyed by file hash) ----------------------------------------
def _cache_path(file_hash: str) -> Path:
    cache_dir = get_storage_root() / "ocr_cache"
    cache_dir.mkdir(parents=True, exist_ok=True)
    return cache_dir / f"{file_hash}.json"


def load_cached_ocr(file_hash: str) -> OcrResult | None:
    path = _cache_path(file_hash)
    if not path.exists():
        return None
    return OcrResult.from_cache(json.loads(path.read_text()))


def save_cached_ocr(file_hash: str, result: OcrResult) -> None:
    _cache_path(file_hash).write_text(json.dumps(result.to_cache()))


# --- orchestration ---------------------------------------------------------
def ocr_document(mime_type: str, storage_path: str) -> OcrResult:
    """Route to the right engine and return a normalized OcrResult (no DB)."""
    if mime_type == PDF_MIME:
        page_texts, page_count, has_layer = probe_pdf_text_layer(storage_path)
        engine = choose_engine(mime_type, has_layer)
        if engine == ENGINE_PDF_TEXT:
            full_text = "\n\n".join(page_texts).strip()
            return OcrResult(
                engine=ENGINE_PDF_TEXT,
                page_count=page_count,
                language=detect_language(full_text),
                text=full_text,
                pages=[OcrPage(page=i, text=t) for i, t in enumerate(page_texts, start=1)],
            )
        return ocr_pdf_via_vision(storage_path)

    if mime_type == DOCX_MIME:
        text = extract_docx(storage_path)
        return OcrResult(
            engine=ENGINE_DOCX,
            page_count=1,
            language=detect_language(text),
            text=text,
            pages=[OcrPage(page=1, text=text)],
        )

    if mime_type in IMAGE_MIMES:
        return ocr_image_via_vision(storage_path)

    raise ValueError(f"Unsupported MIME type for OCR: {mime_type!r}")


def run_ocr(document_id: uuid.UUID, account_id: uuid.UUID) -> None:
    """Background entry point: OCR one document and advance it to `ocr_done`.

    Opens its own session (the request session is gone by now). On success the
    document gains ocr_text/ocr_engine/page_count/language and status `ocr_done`;
    on failure status becomes `failed`. Both outcomes append a processing event.
    """
    started = time.monotonic()
    with SessionLocal() as db:
        document = db.get(Document, document_id)
        if document is None or document.account_id != account_id:
            return  # deleted or wrong account — nothing to do, never cross-scope

        record_event(
            db, account_id=account_id, document_id=document_id,
            stage="ocr", status="started",
        )
        db.commit()

        try:
            cached = load_cached_ocr(document.file_hash)
            cache_hit = cached is not None
            result = cached or ocr_document(document.mime_type, document.storage_path)
            if not cache_hit:
                save_cached_ocr(document.file_hash, result)

            document.ocr_text = result.text
            document.ocr_engine = result.engine
            document.page_count = result.page_count
            document.language = result.language
            document.status = "ocr_done"

            duration_ms = int((time.monotonic() - started) * 1000)
            record_event(
                db, account_id=account_id, document_id=document_id,
                stage="ocr", status="succeeded", duration_ms=duration_ms,
                detail={
                    "engine": result.engine,
                    "page_count": result.page_count,
                    "language": result.language,
                    "char_count": len(result.text),
                    "cache": "hit" if cache_hit else "miss",
                    "failed_pages": result.failed_pages,
                },
            )
            db.commit()

            # Chain extraction now that the document is at `ocr_done`. Local
            # import avoids an import cycle (extraction imports this module for
            # bbox provenance). run_extraction opens its own session and handles
            # its own failures, so it never disturbs the committed OCR result.
            from app.services import extraction

            extraction.run_extraction(document_id, account_id)
        except Exception as exc:  # noqa: BLE001 — record any failure, don't crash the worker
            db.rollback()
            document = db.get(Document, document_id)
            if document is not None:
                document.status = "failed"
                document.error = f"OCR failed: {exc}"
            record_event(
                db, account_id=account_id, document_id=document_id,
                stage="ocr", status="failed",
                error=str(exc),
                duration_ms=int((time.monotonic() - started) * 1000),
            )
            db.commit()
